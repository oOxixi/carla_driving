"""Generate the two-group optimization assignment DOCX."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 90, 90)

GROUP_ONE = [
    "P0｜先冻结评分口径、当前 320 条四模态基线和逐阶段延迟埋点；后续优化统一使用同一验证集对比。",
    "P0｜将标准文本指令从 250 条扩充到不少于 5,000 条；覆盖同义改写、否定、组合指令、模糊目标、目标不存在、安全冲突、单位换算、密集目标、遮挡、曝光和检测误差。",
    "P0｜在第 2 项完成后，按模板、场景和 seed 隔离划分 train/val/test，比例 70%/15%/15%；建立重复、标签冲突和跨 split 泄漏自动检查。",
    "P0｜采集不少于 2,000 条真人音频，覆盖普通话、不同性别/语速/距离、干净环境和经声级计校准的 50 dBA 噪声；方言作为低优先级增益项。",
    "P0｜在第 2—4 项完成后，建立 SenseVoice 文本、意图、槽位、低置信度、安全拒绝及逐语言/噪声回归；门槛为基础意图与槽位不低于 98%，安全拒绝 100%。",
    "P0｜优化语音快路径：流式 VAD、模型常驻预热、SenseVoice INT8/ONNX 或 TensorRT、增量解码、仅高风险低置信度触发第二模型；ASR+NLU P95 目标不高于 60 ms。",
    "P0｜在固定验证集对比 Qwen2.5-VL-7B、3B、AWQ/INT4 版本；按安全正确率、目标关联、显存和 P95 选择模型，不允许只按速度选择。",
    "P0｜在第 7 项选型后优化 Qwen：固定短 JSON、32—48 输出 token、道路 ROI、低像素预算、Top-K 目标、精简场景 JSON、FlashAttention/TensorRT-LLM；复杂指令 P95 目标不高于 300 ms。",
    "P0｜建立 150 ms 快路径：标准速度/启停/减速/明确目标指令由 ASR+确定性 NLU+目标绑定直接下发；Qwen 仅处理复杂、模糊和冲突指令，并保持异步。",
    "P1｜使用扩充数据进行小模型蒸馏或 LoRA 微调；每轮必须同时跑未见模板测试集、检测误差集和安全冲突集，禁止只提高训练集分数。",
    "P1｜输出模型卡、许可证、权重 SHA-256、量化配置、数据版本、逐类准确率及 mean/P95/P99/max 延迟报告。",
]

GROUP_TWO = [
    "P0｜先建立全链路统一时间戳：音频开始、VAD 结束、ASR、NLU、感知、Qwen、规划、安全仲裁和控制生效；输出每阶段 mean/P95/P99/max。",
    "P0｜将控制环与模型推理解耦：控制/安全环保持 20—50 Hz，传感器、ASR 和 Qwen 使用有界队列；任何超时、过期或队列堆积必须 fail-closed。",
    "P0｜优化 RGB 感知：固定道路 ROI、输入缩放、检测低频运行+高频跟踪、Top-K 候选、批处理和 GPU 预处理；感知 P95 目标不高于 30 ms。",
    "P0｜完善毫米波雷达+LiDAR+RGB 融合：统一时间同步、坐标转换、目标关联、TTC/距离估计、单传感器失效降级和误检/漏检注入。",
    "P0｜在第一组第 9 项完成后接入快路径；标准指令端到端 P95 目标不高于 150 ms，安全仲裁与控制合计 P95 目标不高于 5 ms。",
    "P0｜接入异步 Qwen 慢路径：事件触发、结果缓存、deadline、stale/timeout/invalid 拒绝、确定性目标纠正；Qwen 永远不得直接输出油门、刹车或方向盘。",
    "P0｜为 S01、D03、D08 各跑不少于 5 个 seed×20 次；记录成功率、碰撞、闯红灯、最小车距、路线偏差、快路径命中率和端到端延迟。",
    "P1｜扩展密集交通、行人遮挡、夜间、雨雾、强曝光、运动模糊、雷达噪声、传感器掉线、ASR 低置信度和 Qwen 超时场景回归。",
    "P1｜完成不少于 60 分钟稳定性实跑、CARLA/Qwen/ASR 崩溃恢复、显存与队列监控；恢复后自动跑传感器探针和安全制动验证。",
    "P1｜在两组指标冻结后统一生成 evidence_index.json、原始日志哈希、三场景视频、架构图、接口说明、复现脚本、提交检查清单和最终演示材料。",
]


def _set_run(run, *, size: float, bold: bool = False, color=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_numbering(document: Document, num_id: int) -> None:
    numbering = document.part.numbering_part.element
    abstract_id = num_id
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def _apply_number(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def _add_group(document: Document, title: str, tasks: list[str], num_id: int) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    _set_run(heading.add_run(title), size=16, bold=True, color=BLUE)
    for task in tasks:
        paragraph = document.add_paragraph(style="Task")
        _apply_number(paragraph, num_id)
        _set_run(paragraph.add_run(task), size=10.5)


def build(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading = document.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(7)

    task_style = document.styles.add_style("Task", WD_STYLE_TYPE.PARAGRAPH)
    task_style.font.name = FONT
    task_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    task_style.font.size = Pt(10.5)
    task_style.paragraph_format.left_indent = Inches(0.375)
    task_style.paragraph_format.first_line_indent = Inches(-0.194)
    task_style.paragraph_format.space_after = Pt(4)
    task_style.paragraph_format.line_spacing = 1.15
    task_style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run(
        header.add_run("CARLA 挑战赛｜接近满分优化分工"),
        size=8.5,
        color=GRAY,
    )

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(3)
    _set_run(
        title.add_run("两组优化任务分工"),
        size=23,
        bold=True,
        color=RGBColor(0, 0, 0),
    )
    _add_numbering(document, 41)
    _add_numbering(document, 42)
    _add_group(
        document,
        "第一组｜语音、Qwen 与训练数据",
        GROUP_ONE,
        41,
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    _add_group(
        document,
        "第二组｜传感器、控制与全链路",
        GROUP_TWO,
        42,
    )

    for sec in document.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run(footer.add_run("2026-07-28"), size=8.5, color=GRAY)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
